from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn


OUTPUT_PATH = "docs/defense_speech_sprint8.docx"


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 22, "1F4E79"),
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12, "222222"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def main() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Речь для защиты CryptoSafe Manager")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("1F4E79")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Простой сценарий демо-показа с техническими пояснениями").italic = True

    doc.add_paragraph()
    doc.add_heading("Короткое вступление", level=1)
    add_paragraph(
        doc,
        "Здравствуйте. Я показываю CryptoSafe Manager — локальный менеджер паролей. "
        "Главная идея проекта: пользователь хранит пароли локально в своем vault-файле, "
        "доступ защищен мастер-паролем, записи шифруются, важные действия пишутся в аудит, "
        "а импорт, экспорт и обмен данными сделаны через защищенные сценарии.",
    )
    add_paragraph(
        doc,
        "Я буду показывать приложение как обычный пользователь, но параллельно пояснять, "
        "какие модули и механизмы работают внутри.",
    )

    doc.add_heading("Порядок демо", level=1)
    add_numbered(
        doc,
        [
            "Запустить приложение с флешки или из собранной папки.",
            "Выбрать или создать vault-файл.",
            "Пройти мастер первичной настройки и задать мастер-пароль.",
            "Войти в vault.",
            "Создать запись и показать, что пароль скрыт.",
            "Показать генератор паролей.",
            "Показать поиск, категории и теги.",
            "Скопировать пароль и объяснить clipboard-защиту.",
            "Заблокировать и разблокировать vault.",
            "Открыть журнал аудита и проверить целостность.",
            "Показать защищенный экспорт и импорт.",
            "Показать Bitwarden export/import, если потребуется.",
            "Показать Share одной записи и QR/key exchange.",
            "Показать Panic mode и tray.",
            "В конце сказать про тесты, coverage, документацию и PyInstaller-сборку.",
        ],
    )

    doc.add_heading("1. Установка и запуск", level=1)
    add_paragraph(
        doc,
        "Что показываю: запускаю CryptoSafe Manager из собранной папки.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: приложение упаковано через PyInstaller. Поэтому на другом компьютере "
        "его можно запускать как desktop-приложение, а не как набор Python-файлов.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: точка входа проекта — run.py. Сборка описана в cryptosafe-manager.spec. "
        "Результат сборки лежит в dist/CryptoSafe Manager, а для передачи используется ZIP всего каталога.",
        "Технически:",
    )

    doc.add_heading("2. Выбор vault и создание базы", level=1)
    add_paragraph(
        doc,
        "Что показываю: выбираю существующий vault или создаю новый файл базы.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: сначала выбирается vault-файл, и только потом вводится мастер-пароль. "
        "Это сделано потому, что у пользователя может быть несколько разных vault.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: vault хранится в SQLite. За базу отвечает src/database/db.py. "
        "Схема базы мигрируется через PRAGMA user_version, поэтому старые версии можно обновлять.",
        "Технически:",
    )

    doc.add_heading("3. Мастер-пароль и вход", level=1)
    add_paragraph(
        doc,
        "Что показываю: задаю мастер-пароль в мастере настройки и вхожу в vault.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: мастер-пароль не хранится в базе в открытом виде. Он нужен для проверки входа "
        "и получения ключа, которым дальше шифруются данные.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: проверка мастер-пароля выполняется через Argon2id. Ключ шифрования получается "
        "через PBKDF2-HMAC-SHA256. Этим занимаются AuthenticationService, KeyDerivation и KeyStorage.",
        "Технически:",
    )

    doc.add_heading("4. Создание записи и шифрование", level=1)
    add_paragraph(
        doc,
        "Что показываю: создаю запись с названием, логином, паролем, категорией и тегами.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: пароль и чувствительные данные записи не лежат в базе plaintext-ом. "
        "Они шифруются перед сохранением и расшифровываются только после входа в vault.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: новые записи шифруются через AES-256-GCM. Это важно, потому что GCM дает не только "
        "шифрование, но и проверку целостности: если ciphertext изменить, расшифровать его уже не получится.",
        "Технически:",
    )
    add_paragraph(
        doc,
        "Где в коде: src/core/vault/encryption_service.py и src/core/vault/entry_manager.py.",
        "Где в коде:",
    )

    doc.add_heading("5. Генератор паролей", level=1)
    add_paragraph(
        doc,
        "Что показываю: генерирую пароль в окне добавления или изменения записи.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: пользователь может не придумывать пароль вручную. Приложение предлагает защищенные "
        "параметры по умолчанию: длину, разные типы символов и проверку сложности.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: генератор находится в src/core/vault/password_generator.py, а оценка сложности — "
        "в src/core/crypto/password_validator.py.",
        "Технически:",
    )

    doc.add_heading("6. Поиск, категории и теги", level=1)
    add_paragraph(
        doc,
        "Что показываю: нахожу запись через поиск, фильтрую по категории или тегу.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: это слой удобства для пользователя. Он помогает быстро работать с большим vault, "
        "но не отменяет шифрование чувствительных данных.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: поиск и фильтрация проходят через EntryManager и SearchIndex.",
        "Технически:",
    )

    doc.add_heading("7. Clipboard", level=1)
    add_paragraph(
        doc,
        "Что показываю: копирую пароль и показываю уведомление/статус clipboard.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: clipboard — опасное место, потому что системный буфер могут читать другие приложения. "
        "Поэтому копирование вынесено в отдельный сервис.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: ClipboardService ставит таймер auto-clear, очищает буфер, отслеживает внешние изменения, "
        "публикует события и поддерживает memory-only режим. PlatformAdapter отвечает за Windows, macOS и Linux.",
        "Технически:",
    )
    add_paragraph(
        doc,
        "Где в коде: src/core/clipboard/clipboard_service.py и src/core/clipboard/platform_adapter.py.",
        "Где в коде:",
    )

    doc.add_heading("8. Lock и Unlock", level=1)
    add_paragraph(
        doc,
        "Что показываю: нажимаю Lock, затем Unlock.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: при блокировке приложение очищает активное состояние и требует повторный мастер-пароль. "
        "Это нужно, чтобы данные не оставались доступными после отхода пользователя от компьютера.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: StateManager хранит состояние сессии, таймеры неактивности и активный ключ. "
        "При блокировке ключ и видимые секреты очищаются.",
        "Технически:",
    )

    doc.add_heading("9. Журнал аудита", level=1)
    add_paragraph(
        doc,
        "Что показываю: открываю журнал аудита, обновляю список и запускаю проверку целостности.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: аудит фиксирует важные события: вход, выход, добавление записей, clipboard, импорт, экспорт, "
        "share и panic mode. Это нужно, чтобы пользователь видел историю действий.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: аудит защищен hash chain и подписью. Если изменить или удалить старую запись, проверка "
        "целостности должна это обнаружить.",
        "Технически:",
    )
    add_paragraph(
        doc,
        "Где в коде: src/core/audit/audit_logger.py, log_signer.py и log_verifier.py.",
        "Где в коде:",
    )

    doc.add_heading("10. Импорт и экспорт", level=1)
    add_paragraph(
        doc,
        "Что показываю: делаю encrypted export, затем import этого файла.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: основной экспорт защищенный. Это не plaintext-файл с паролями, а encrypted JSON. "
        "Его можно использовать как резервную копию или для переноса.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: в export-файле есть ciphertext, salt, nonce, checksum/HMAC и metadata. "
        "При импорте файл валидируется, очищается от потенциально опасного содержимого и только потом данные пишутся в vault.",
        "Технически:",
    )
    add_paragraph(
        doc,
        "Где в коде: src/core/import_export/exporter.py и importer.py.",
        "Где в коде:",
    )

    doc.add_heading("11. Bitwarden и совместимость", level=1)
    add_paragraph(
        doc,
        "Что показываю: если попросят, показываю экспорт/импорт для Bitwarden.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: я специально не делаю основной сценарий через plaintext CSV, потому что это оставляет пароли "
        "на диске в открытом виде. Для совместимости используется защищенный JSON-сценарий.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: форматы password manager находятся в src/core/import_export/formats/password_manager.py.",
        "Технически:",
    )

    doc.add_heading("12. Secure sharing и QR/key exchange", level=1)
    add_paragraph(
        doc,
        "Что показываю: выбираю одну запись и открываю Share. Затем показываю QR/key exchange.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: share передает не весь vault, а только одну выбранную запись. Это снижает риск раскрытия остальных данных.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: SharingService создает зашифрованный пакет. Для public-key сценария используется RSA/ECC. "
        "QR/key exchange передает публичный ключ, а private key остается у владельца.",
        "Технически:",
    )
    add_paragraph(
        doc,
        "Где в коде: src/core/import_export/sharing_service.py и key_exchange.py.",
        "Где в коде:",
    )

    doc.add_heading("13. Panic mode и tray", level=1)
    add_paragraph(
        doc,
        "Что показываю: показываю tray-меню и panic mode.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: panic mode нужен для экстренной ситуации. Он блокирует vault, очищает clipboard, скрывает окно "
        "и записывает событие в аудит.",
        "Что говорю:",
    )
    add_paragraph(
        doc,
        "Технически: panic mode находится в src/core/security/panic_mode.py, а tray-интеграция подключена в MainWindow через pystray.",
        "Технически:",
    )

    doc.add_heading("14. Тесты и финальный Sprint 8", level=1)
    add_paragraph(
        doc,
        "Что показываю: при необходимости открываю tests/report/summary.md или coverage HTML.",
        "Что показываю:",
    )
    add_paragraph(
        doc,
        "Что говорю: в финальном спринте подготовлен pytest-набор. Основной набор проходит меньше чем за 30 секунд, "
        "coverage выше 80 процентов, а медленные stress и memory dump тесты вынесены отдельно.",
        "Что говорю:",
    )
    add_bullets(
        doc,
        [
            "Основной запуск: python -m pytest.",
            "Coverage: python -m pytest --cov=src.",
            "Slow tests: python -m pytest -m slow.",
            "Отчет: tests/report/summary.md и tests/report/html/index.html.",
            "Integrity-тест: tests/test_project_integrity.py проверяет циклические импорты и технические хвосты.",
        ],
    )

    doc.add_heading("Короткие ответы на вопросы", level=1)
    questions = [
        ("Почему мастер-пароль не хранится?", "Потому что хранится только результат проверки и параметры derivation. Сам пароль нужен только во время входа."),
        ("Почему AES-GCM?", "Потому что он дает шифрование и проверку целостности. Если ciphertext подменить, расшифровка не пройдет."),
        ("Где plaintext?", "Только кратковременно в памяти после расшифровки или при копировании. При lock состояние очищается."),
        ("Почему clipboard отдельно?", "Потому что системный буфер опасен. Его могут читать другие приложения, поэтому нужен auto-clear и мониторинг."),
        ("Как защищен аудит?", "Через sequence number, hash chain и подпись записей."),
        ("Зачем slow-тесты отдельно?", "Чтобы основной тестовый набор укладывался в лимит времени, но тяжелые проверки не удалялись."),
        ("Что делать, если tray не работает?", "Tray зависит от ОС и backend. Приложение должно безопасно работать и без tray через основное окно."),
    ]
    for question, answer in questions:
        add_paragraph(doc, f"Вопрос: {question}", "Вопрос:")
        add_paragraph(doc, f"Ответ: {answer}", "Ответ:")

    doc.add_heading("Финальная фраза", level=1)
    add_paragraph(
        doc,
        "В итоге проект доведен до состояния законченного desktop-приложения: есть установка и запуск, выбор vault, "
        "мастер-пароль, шифрование записей, безопасный clipboard, аудит, импорт и экспорт, secure sharing, QR/key exchange, "
        "panic mode, тесты, coverage, документация и PyInstaller-сборка.",
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
